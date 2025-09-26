import streamlit as st
import google.generativeai as genai
import docx
import fitz  # PyMuPDF
import io
import re
from docx.enum.text import WD_COLOR_INDEX
import pandas as pd
from docx.shared import Pt
import zipfile

# --- Konfigurasi Awal Halaman ---
st.set_page_config(
    page_title="Proofreader Bahasa Indonesia",
    layout="wide"
)

# --- Header Aplikasi ---
try:
    # Ganti "Logo_IFG-removebg-preview.png" dengan nama file logo Anda
    st.image("Logo_IFG-removebg-preview.png", width=250)
except Exception:
    st.warning("Logo tidak ditemukan. Pastikan file logo ada di direktori yang sama.")

st.markdown("<h1 style='text-align: center;'>Website Proofreader COE Divisi SKAI IFG</h1>", unsafe_allow_html=True)

# --- BAGIAN 1: PROOFREAD ---
st.divider() # Garis pemisah
st.header("1. Proofread Dokumen")

# --- Inisialisasi Session State ---
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None

# --- Konfigurasi API Key Google ---
try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-pro-latest')
except KeyError:
    st.error("Google API Key belum diatur. Harap atur di Streamlit Secrets.")
    st.stop()
except Exception as e:
    st.error(f"Terjadi masalah saat mengkonfigurasi Google AI: {e}")
    st.stop()

# --- FUNGSI-FUNGSI UTAMA ---

def extract_text_with_pages(uploaded_file):
    """Mengekstrak teks dari file PDF atau DOCX."""
    pages_content = []
    file_extension = uploaded_file.name.split('.')[-1].lower()
    file_bytes = uploaded_file.getvalue()

    if file_extension == 'pdf':
        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num, page in enumerate(pdf_document):
                pages_content.append({"halaman": page_num + 1, "teks": page.get_text()})
            pdf_document.close()
        except Exception as e:
            st.error(f"Gagal membaca file PDF: {e}")
            return None
    elif file_extension == 'docx':
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            pages_content.append({"halaman": 1, "teks": full_text})
        except Exception as e:
            st.error(f"Gagal membaca file DOCX: {e}")
            return None
    else:
        st.error("Format file tidak didukung. Harap unggah .pdf atau .docx")
        return None
    return pages_content

def proofread_with_gemini(text_to_check):
    """Mengirim teks ke Gemini untuk proofreading dan mem-parsing hasilnya."""
    if not text_to_check or text_to_check.isspace():
        return []
    
    prompt = f"""
    Anda adalah seorang auditor dan ahli bahasa Indonesia yang sangat teliti.
    Tugas Anda adalah melakukan proofread pada teks berikut. Fokus pada:
    1. Memperbaiki kesalahan ketik (typo).
    2. Memastikan semua kata sesuai dengan Kamus Besar Bahasa Indonesia (KBBI).
    3. Memperbaiki kesalahan tata bahasa sederhana dan ejaan agar sesuai dengan Pedoman Umum Ejaan Bahasa Indonesia (PUEBI).
    4. Jika ada yang bahasa inggris, tolong di italic
    5. Nama-nama yang diberi ini pastikan benar juga "Yullyan, I Made Suandi Putra, Laila Fajriani, Hari Sundoro, Bakhas Nasrani Diso, Rizky Ananda Putra, Wirawan Arief Nugroho, Lelya Novita Kusumawati, Ryani Ariesti Syafitri, Darmo Saputro Wibowo, Lucky Parwitasari, Handarudigdaya Jalanidhi Kuncaratrah, Fajar Setianto, Jaka Tirtana Hanafiah,  Muhammad Rosyid Ridho Muttaqien, Octovian Abrianto, Deny Sjahbani, Jihan Abigail, Winda Anggraini, Fadian Dwiantara, Aliya Anindhita Rachman"
    6. Fontnya arial dan jangan diganti. Khusus untuk judul paling atas, itu font sizenya 12 dan bodynya selalu 11
    7. Khusus "Indonesia Financial Group (IFG)", meskipun bahasa inggris, tidak perlu di italic
    8. Bila ada bahasa yang lebih bagus, tolong berikan saran dan diberi warna highlight yang berbeda selain kuning
    9. Pada bagian penutup, yang mulai dari "Jakarta, 4 September 2025" hingga "Kepala Divisi Satuan Kerja Audit Internal", tidak perlu dicek
    10. Pada bagian nomor surat, itu juga tidak perlu di cek
    11. Kalau ada kata-kata yang tidak sesuai KBBI dan PUEBI, tolong jangan highlight semua kalimatnya, tapi cukup highlight kata-kata yang salah serta perbaiki kata-kata itu aja, jangan perbaiki semua kalimatnya
    12. Ketika Anda perbaiki, fontnya pastikan Arial dengan ukuran 11 juga
    13. Kalau ada kata-kata dalam bahasa inggris, tolong jangan Anda sarankan untuk ganti ke bahasa indonesia, cukup Anda sarankan untuk italic
    14. Pada kalimat "Indonesia Financial Group", jika terdapat kata typo "Finansial", tolong Anda sarankan untuk ganti ke "Financial"
    15. Yang benar adalah "Satuan Kerja Audit Internal", bukan "Satuan Pengendali Internal Audit"
    16. Jika terdapat kata "reviu", biarkan itu sebagai benar
    17. Kalau ada kata "IM", "ST", "SKAI", "IFG", "TV (Angka Romawi)", "RKAT", dan "RKAP", itu tidak perlu ditandai sebagai salah dan tidak perlu disarankan untuk italic / bold / underline
    18. Kalau ada kata "email", biarkan itu sebagai benar tapi disarankan italic saja
    19. Untuk nama modul seperti "Modul Sourcing, dll", itu tidak perlu italic
    20. Kalau ada kata dalam bahasa inggris yang masih masuk akal dan nyambung dengan kalimat yang dibahas, tidak perlu Anda sarankan untuk ganti ke bahasa indonesia
    21. Jika ada bahasa inggris dan akronimnya seperti "General Ledger (GL)", tolong dilakukan italic pada kata tersebut pada saat download file hasil revisinya, akronimnya tidak usah\
    22. Awal kalimat selalu dimulai dengan huruf kapital

    PENTING: Berikan hasil dalam format yang SANGAT KETAT. Untuk setiap kesalahan, gunakan format:
    [SALAH] kata atau frasa yang salah -> [BENAR] kata atau frasa perbaikan -> [KALIMAT] kalimat lengkap asli tempat kesalahan ditemukan

    Contoh:
    [SALAH] dikarenakan -> [BENAR] karena -> [KALIMAT] Hal itu terjadi dikarenakan kelalaian petugas.

    Jika tidak ada kesalahan sama sekali, kembalikan teks: "TIDAK ADA KESALAHAN"

    Berikut adalah teks yang harus Anda periksa:
    ---
    {text_to_check}
    """
    try:
        response = model.generate_content(prompt)
        pattern = re.compile(r"\[SALAH\]\s*(.*?)\s*->\s*\[BENAR\]\s*(.*?)\s*->\s*\[KALIMAT\]\s*(.*?)\s*(\n|$)", re.IGNORECASE)
        found_errors = pattern.findall(response.text)
        return [{"salah": salah.strip(), "benar": benar.strip(), "kalimat": kalimat.strip()} for salah, benar, kalimat, _ in found_errors]
    except Exception as e:
        st.error(f"Terjadi kesalahan saat menghubungi AI: {e}")
        return []

def generate_revised_docx(file_bytes, errors):
    """
    Membuat dokumen .docx dengan semua kesalahan yang sudah diperbaiki
    SAMBIL MEMPERTAHANKAN FONT ASLI (Arial 11 untuk body dan Arial 12 untuk judul).
    """
    doc = docx.Document(io.BytesIO(file_bytes))

    for error in reversed(errors):
        salah = error["Kata/Frasa Salah"]
        benar = error["Perbaikan Sesuai KBBI"]
        for para in doc.paragraphs:
            if salah in para.text:
                original_font = None
                if para.runs:
                    original_font = para.runs[0].font
                
                current_text = para.text
                para.text = current_text.replace(salah, benar, 1)

                if original_font:
                    for run in para.runs:
                        font = run.font
                        font.name = original_font.name
                        font.size = original_font.size
                        font.bold = original_font.bold
                        font.italic = original_font.italic
                        font.underline = original_font.underline
                        font.color.rgb = original_font.color.rgb

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def generate_highlighted_docx(file_bytes, errors):
    """
    Membuat dokumen .docx dengan semua kesalahan yang di-highlight.
    Versi ini sudah diperbaiki untuk menangani kata yang terpecah di antara 'runs'.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    unique_salah = set(error["Kata/Frasa Salah"] for error in errors)

    for para in doc.paragraphs:
        # Cek setiap kata yang salah untuk setiap paragraf
        for term in unique_salah:
            # Lakukan pencarian case-insensitive (tidak membedakan huruf besar/kecil)
            if term.lower() in para.text.lower():
                # Jika kata ditemukan, kita akan membangun ulang paragraf ini
                full_text = para.text
                
                # Simpan gaya (font, ukuran, dll.) dari run pertama sebagai dasar
                # Catatan: Ini akan membuat seluruh paragraf memiliki gaya yang sama
                original_style_run = para.runs[0] if para.runs else None
                
                # Kosongkan paragraf yang ada
                para.clear()

                # Gunakan regex untuk memecah teks berdasarkan kata yang salah, sambil mempertahankannya
                # Ini akan menemukan semua kemunculan kata, tidak peduli besar kecilnya huruf
                parts = re.split(f'({re.escape(term)})', full_text, flags=re.IGNORECASE)

                for part in parts:
                    if part: # Pastikan bagian tidak kosong
                        # Jika bagian ini sama dengan kata yang kita cari (case-insensitive)
                        if part.lower() == term.lower():
                            run = para.add_run(part)
                            # Beri highlight
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            # Jika tidak, tambahkan sebagai teks biasa
                            run = para.add_run(part)
                        
                        # Terapkan kembali gaya dasar ke setiap run baru
                        if original_style_run:
                            run.font.name = original_style_run.font.name
                            run.font.size = original_style_run.font.size
                            run.bold = original_style_run.bold
                            run.italic = original_style_run.italic

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def create_zip_archive(revised_data, highlighted_data, original_filename):
    """Menggabungkan dua file DOCX ke dalam satu arsip ZIP di memori."""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr(f"revisi_{original_filename}", revised_data)
        zip_file.writestr(f"highlight_{original_filename}", highlighted_data)
    return zip_buffer.getvalue()

# --- ANTARMUKA STREAMLIT UNTUK BAGIAN 1 ---
uploaded_file = st.file_uploader(
    "Unggah dokumen (DOCX)",
    type=['docx'],
    help="File yang diunggah akan dianalisis untuk menemukan kesalahan ejaan dan ketik."
)

if uploaded_file is not None:
    st.info(f"File yang diunggah: **{uploaded_file.name}**")

    if st.button("Mulai Analisis", type="primary", use_container_width=True):
        with st.spinner("Membaca dan menganalisis dokumen..."):
            document_pages = extract_text_with_pages(uploaded_file)
            
            if document_pages:
                all_errors = []
                progress_bar = st.progress(0, text="Menganalisis teks dengan AI...")
                
                for i, page in enumerate(document_pages):
                    progress_text = f"Menganalisis Bagian {i + 1}/{len(document_pages)}..."
                    progress_bar.progress((i + 1) / len(document_pages), text=progress_text)
                    found_errors_on_page = proofread_with_gemini(page['teks'])
                    
                    for error in found_errors_on_page:
                        all_errors.append({
                            "Kata/Frasa Salah": error['salah'],
                            "Perbaikan Sesuai KBBI": error['benar'],
                            "Pada Kalimat": error['kalimat'],
                            "Ditemukan di Halaman": page['halaman']
                        })
                progress_bar.empty()
                st.session_state.analysis_results = all_errors

if st.session_state.analysis_results is not None:
    all_errors = st.session_state.analysis_results

    if not all_errors:
        st.success("Tidak ada kesalahan ejaan atau ketik yang ditemukan dalam dokumen.")
    else:
        st.warning(f"Ditemukan **{len(all_errors)}** potensi kesalahan dalam dokumen.")
        
        df_errors = pd.DataFrame(all_errors)
        st.dataframe(df_errors, use_container_width=True)
        
        st.subheader("Unduh Hasil")
        
        if uploaded_file and uploaded_file.name.endswith('.docx'):
            with st.spinner("Mempersiapkan semua file unduhan..."):
                revised_docx_data = generate_revised_docx(uploaded_file.getvalue(), all_errors)
                highlighted_docx_data = generate_highlighted_docx(uploaded_file.getvalue(), all_errors)

            col1, col2, col3 = st.columns(3)

            with col1:
                st.download_button(
                    label="Unduh Direvisi (.docx)",
                    data=revised_docx_data,
                    file_name=f"revisi_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col2:
                st.download_button(
                    label="Unduh Highlight (.docx)",
                    data=highlighted_docx_data,
                    file_name=f"highlight_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col3:
                zip_data = create_zip_archive(revised_docx_data, highlighted_docx_data, uploaded_file.name)
                st.download_button(
                    label="Unduh Semua (.zip)",
                    data=zip_data,
                    file_name=f"hasil_proofread_{uploaded_file.name.split('.')[0]}.zip",
                    mime="application/zip",
                    use_container_width=True
                )
    st.warning("Harap lakukan analisis ulang jika muncul pesan bahwa tidak ada ejaan yang salah karena model bisa saja tidak berhasil mendeteksi.")


st.divider()
st.header("2. Bandingkan Dokumen")

# Tambahan import yang mungkin diperlukan untuk bagian ini
import difflib
from docx import Document
from docx.shared import Pt, Inches

# --- Fungsi-fungsi Helper untuk Bagian 2 ---

def extract_paragraphs(docx_file):
    """Membaca file DOCX yang diunggah dan mengembalikan isinya sebagai daftar paragraf."""
    try:
        source_stream = io.BytesIO(docx_file.getvalue())
        doc = docx.Document(source_stream)
        return [p.text for p in doc.paragraphs if p.text.strip() != ""]
    except Exception as e:
        st.error(f"Gagal membaca file {docx_file.name}: {e}")
        return []

def get_revision_confidence(original_sentence, revised_sentence):
    """Meminta model AI untuk memberikan skor keyakinan bahwa revisi lebih baik."""
    if original_sentence == revised_sentence:
        return 100
    prompt = f"""
    Bandingkan dua kalimat ini berdasarkan PUEBI dan KBBI (Pastikan Anda review PUEBI dan KBBI terlebih dahulu sebelum bandingkan).
    Kalimat Asli: "{original_sentence}"
    Kalimat Revisi: "{revised_sentence}"

    Beri skor keyakinan (angka 0-100) bahwa revisi tersebut adalah perbaikan yang benar dan diperlukan.
    PENTING: Jawab HANYA dengan ANGKA.
    """
    try:
        response = model.generate_content(prompt)
        confidence_str = ''.join(filter(str.isdigit, response.text))
        return int(confidence_str) if confidence_str else "N/A"
    except Exception:
        return "N/A"

def find_word_diff(original_para, revised_para):
    """Menemukan dan menyorot kata-kata yang berbeda antara dua paragraf."""
    matcher = difflib.SequenceMatcher(None, original_para.split(), revised_para.split())
    diffs = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace' or tag == 'insert':
            diffs.append(" ".join(revised_para.split()[j1:j2]))
    return ", ".join(diffs) if diffs else "Perubahan Minor"

def create_comparison_docx(df):
    """
    Membuat file DOCX dari DataFrame hasil perbandingan
    dengan highlight pada kata yang direvisi.
    """
    doc = Document()
    # Menambahkan Judul Utama Dokumen
    title = doc.add_heading('Hasil Perbandingan Dokumen', level=1)
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True
    doc.add_paragraph()

    # Menambahkan Tabel
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Menambahkan Header Tabel (Arial 11 Bold)
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        p = hdr_cells[i].paragraphs[0]
        p.text = col_name
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.size = Pt(11)

    # Menambahkan Isi Tabel dengan Logika Highlighting
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        # Iterasi melalui setiap sel di baris
        for i, col_name in enumerate(df.columns):
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_value = str(row[col_name])
            
            # --- LOGIKA BARU UNTUK HIGHLIGHTING ---
            # Jika ini adalah kolom 'Kalimat Revisi', terapkan highlight
            if col_name == "Kalimat Revisi":
                original_text = str(row["Kalimat Awal"])
                revised_text = cell_value
                
                # Gunakan difflib untuk menemukan perbedaan kata per kata
                matcher = difflib.SequenceMatcher(None, original_text.split(), revised_text.split())
                
                for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                    # Ambil bagian dari kalimat revisi
                    text_segment = " ".join(revised_text.split()[j1:j2]) + " "
                    
                    if tag == 'equal':
                        # Jika sama, tambahkan sebagai teks biasa
                        run = cell_paragraph.add_run(text_segment)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                    else: # Untuk 'replace' atau 'insert', beri highlight
                        run = cell_paragraph.add_run(text_segment)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                # Untuk kolom lain, tulis teks seperti biasa
                run = cell_paragraph.add_run(cell_value)
                run.font.name = 'Arial'
                run.font.size = Pt(11)

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

# --- Antarmuka Streamlit untuk Bagian 2 ---

col1, col2 = st.columns(2)
with col1:
    original_file = st.file_uploader("Unggah Dokumen Asli", type=['docx'], key="original_doc")
with col2:
    proofread_file = st.file_uploader("Unggah Dokumen Hasil Proofread", type=['docx'], key="proofread_doc")

if original_file is not None and proofread_file is not None:
    if st.button("Bandingkan Dokumen", use_container_width=True, type="primary"):
        with st.spinner("Mengekstrak teks dan membandingkan dokumen..."):
            original_paras = extract_paragraphs(original_file)
            revised_paras = extract_paragraphs(proofread_file)
            comparison_results = []
            matcher = difflib.SequenceMatcher(None, original_paras, revised_paras)

            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'replace':
                    for i in range(i1, i2):
                        original_para = original_paras[i]
                        revised_para = revised_paras[j1 + (i - i1)] if (j1 + (i - i1)) < j2 else ""
                        if revised_para:
                            word_diff = find_word_diff(original_para, revised_para)
                            confidence = get_revision_confidence(original_para, revised_para)
                            comparison_results.append({
                                "Kalimat Awal": original_para,
                                "Kalimat Revisi": revised_para,
                                "Kata yang Direvisi": word_diff,
                            })
            
            # Simpan hasil perbandingan ke session state
            st.session_state.comparison_results = pd.DataFrame(comparison_results)

# Menampilkan hasil jika ada di session state
if 'comparison_results' in st.session_state and not st.session_state.comparison_results.empty:
    df_comparison = st.session_state.comparison_results
    st.success(f"Perbandingan selesai. Ditemukan {len(df_comparison)} paragraf yang direvisi.")
    st.dataframe(df_comparison, use_container_width=True)

    # Menambahkan tombol download untuk hasil perbandingan
    docx_data = create_comparison_docx(df_comparison)
    st.download_button(
        label="Unduh Hasil Perbandingan (.docx)",
        data=docx_data,
        file_name=f"perbandingan_{original_file.name}",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

elif 'comparison_results' in st.session_state:
     st.info("Tidak ditemukan perbedaan signifikan antar paragraf di kedua dokumen.")

# --- BAGIAN 3: ANALISIS KOHERENSI DOKUMEN ---
st.divider()
st.header("3. Analisis Koherensi Dokumen")

def analyze_document_coherence(full_text):
    """Mengirim teks ke AI untuk dianalisis koherensinya dan memberikan saran."""
    if not full_text or full_text.isspace():
        return []

    prompt = f"""
    Anda adalah seorang editor ahli yang bertugas menganalisis struktur dan koherensi sebuah tulisan.
    Tugas Anda adalah membaca keseluruhan teks berikut dan mengidentifikasi setiap kalimat atau paragraf yang tidak koheren atau keluar dari topik utama di dalam sebuah sub-bagian.
    
    Untuk setiap ketidaksesuaian yang Anda temukan, lakukan hal berikut:
    1. Bacalah mengenai judul dari section atau subsection yang ada pada file tersebut
    2. Tentukan topik utama dari paragraf atau bagian tempat kalimat itu berada.
    3. Identifikasi kalimat asli yang menyimpang dari topik tersebut.
    4. Bila ada kalimat yang sekiranya memyimpang, Berikan saran dengan menghighlight kalimat tersebut untuk diulis ulang kalimat tersebut agar relevan dan menyatu kembali dengan topik utamanya, sambil berusaha mempertahankan maksud aslinya jika memungkinkan.

    Berikan hasil dalam format yang SANGAT KETAT seperti di bawah ini. Ulangi format ini untuk setiap kalimat menyimpang yang Anda temukan:
    [TOPIK UTAMA] topik utama dari bagian tersebut -> [TEKS ASLI] kalimat asli yang tidak koheren -> [SARAN REVISI] versi kalimat yang sudah diperbaiki agar koheren

    Contoh:
    [TOPIK UTAMA] Sistem Whistleblowing Perusahaan -> [TEKS ASLI] Selain itu, audit internal juga memeriksa laporan keuangan setiap kuartal. -> [SARAN REVISI] Sistem whistleblowing ini terintegrasi dengan audit internal untuk menindaklanjuti laporan yang masuk, terutama yang berkaitan dengan anomali keuangan.

    Jika seluruh dokumen sudah koheren dan tidak ada masalah, kembalikan teks: "TIDAK ADA MASALAH KOHERENSI"

    Berikut adalah teks yang harus dianalisis:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        pattern = re.compile(r"\[TOPIK UTAMA\]\s*(.*?)\s*->\s*\[TEKS ASLI\]\s*(.*?)\s*->\s*\[SARAN REVISI\]\s*(.*?)\s*(\n|$)", re.IGNORECASE)
        found_issues = pattern.findall(response.text)
        return [{"topik": topik.strip(), "asli": asli.strip(), "saran": saran.strip()} for topik, asli, saran, _ in found_issues]
    except Exception as e:
        st.error(f"Terjadi kesalahan saat menghubungi AI: {e}")
        return []

# Antarmuka Streamlit untuk Bagian 3
coherence_file = st.file_uploader(
    "Unggah dokumen untuk dianalisis koherensinya",
    type=['docx'],
    key="coherence_doc"
)

if coherence_file is not None:
    if st.button("Analisis Koherensi", use_container_width=True, type="primary"):
        with st.spinner("Membaca dan menganalisis struktur dokumen..."):
            # Kita gunakan fungsi extract_text_with_pages yang sudah ada
            document_pages = extract_text_with_pages(coherence_file)
            if document_pages:
                # Gabungkan semua teks dari halaman menjadi satu
                full_document_text = "\n".join([page['teks'] for page in document_pages])
                
                # Panggil fungsi analisis yang baru
                coherence_issues = analyze_document_coherence(full_document_text)

                # Simpan hasilnya ke session state
                st.session_state.coherence_results = coherence_issues

# Menampilkan hasil jika ada di session state
if 'coherence_results' in st.session_state:
    results = st.session_state.coherence_results
    if not results:
        st.success("Analisis selesai. Tidak ditemukan masalah koherensi yang signifikan dalam dokumen.")
    else:
        st.warning(f"Analisis selesai. Ditemukan {len(results)} potensi masalah koherensi.")
        
        # Mengubah nama kolom untuk tampilan yang lebih baik
        df_coherence = pd.DataFrame(results)
        df_coherence.rename(columns={
            'topik': 'Topik Utama Seharusnya',
            'asli': 'Teks Asli (Tidak Koheren)',
            'saran': 'Saran Revisi (Koheren)'
        }, inplace=True)

        st.dataframe(df_coherence, use_container_width=True)

st.divider()
st.markdown('### 4. Restrukturisasi Koherensi Dokumen <b style="color:red;">(UNAVAILABLE)</b>', unsafe_allow_html=True)

def get_structural_recommendations(full_text):
    """Meminta AI untuk menganalisis dan memberikan saran pemindahan paragraf."""
    if not full_text or full_text.isspace():
        return []

    prompt = f"""
    Anda adalah seorang editor struktural ahli. Tugas Anda adalah menganalisis draf dokumen berikut untuk menemukan paragraf yang tidak koheren atau "tersesat" (tidak sesuai dengan topik utama sub-babnya).

    Untuk setiap paragraf yang tersesat, Anda harus:
    1.  Bacalah semua dokumennya terlebih dahulu sebelum Anda membuat revisi
    1.  Pada saat Anda baca dokumennya, tolong Identifikasi teks lengkap dari paragraf yang tidak pada tempatnya.
    2.  Tentukan di bab atau sub-bab mana paragraf itu berada saat ini (lokasi asli).
    3.  Berikan rekomendasi di bab atau sub-bab mana paragraf tersebut seharusnya diletakkan agar lebih koheren dan masuk akal.

    Berikan hasil dalam format JSON yang berisi sebuah list. Setiap objek dalam list harus memiliki tiga kunci: "original_location", "misplaced_paragraph", dan "recommended_location".

    Contoh Format JSON:
    [
      {{
        "original_location": "Bab 2.1: Prosedur Whistleblowing",
        "misplaced_paragraph": "Selain itu, audit internal juga bertugas memeriksa laporan keuangan setiap kuartal untuk memastikan tidak ada anomali.",
        "recommended_location": "Bab 4.2: Peran Audit Internal dalam Pengawasan Keuangan"
      }},
      {{
        "original_location": "Bab 1: Pendahuluan",
        "misplaced_paragraph": "Proses rekrutmen karyawan baru akan dimulai bulan depan dengan membuka lowongan di berbagai platform.",
        "recommended_location": "Bab 5: Sumber Daya Manusia dan Rekrutmen"
      }}
    ]

    Jika seluruh dokumen sudah terstruktur dengan baik, kembalikan list kosong: []

    Berikut adalah teks dokumen yang harus dianalisis:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        cleaned_response = re.sub(r'```json\s*|\s*```', '', response.text.strip())
        
        import json
        recommendations = json.loads(cleaned_response)
        return recommendations
    except Exception as e:
        st.error(f"Gagal memproses atau mem-parsing respons dari AI: {e}")
        return []

def create_recommendation_excel(df):
    """Membuat file Excel dari DataFrame hasil rekomendasi."""
    output_buffer = io.BytesIO()
    with pd.ExcelWriter(output_buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Rekomendasi Struktur')
        
        # Mengatur lebar kolom agar mudah dibaca
        worksheet = writer.sheets['Rekomendasi Struktur']
        for i, col in enumerate(df.columns):
            column_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
            worksheet.set_column(i, i, column_len)
            
    return output_buffer.getvalue()

# --- Antarmuka Streamlit untuk Bagian 3 ---
recommendation_file = st.file_uploader(
    "Unggah dokumen untuk mendapatkan saran restrukturisasi",
    type=['pdf', 'docx'],
    key="recommendation_doc"
)

if recommendation_file is not None:
    if st.button("Dapatkan Rekomendasi Struktur", use_container_width=True, type="primary"):
        with st.spinner("Menganalisis keseluruhan struktur dokumen..."):
            document_pages = extract_text_with_pages(recommendation_file)
            if document_pages:
                full_text = "\n".join([page['teks'] for page in document_pages])
                recommendations = get_structural_recommendations(full_text)
                
                # Simpan hasilnya ke session state
                st.session_state.recommendations = recommendations

# Menampilkan hasil dan tombol download jika ada
if 'recommendations' in st.session_state:
    results = st.session_state.recommendations
    
    if not results:
        st.success("Analisis selesai. Struktur dokumen Anda sudah koheren dan tidak ditemukan paragraf yang perlu dipindahkan.")
    else:
        st.warning(f"Analisis selesai. Ditemukan {len(results)} paragraf yang disarankan untuk dipindahkan.")
        
        df_recommendations = pd.DataFrame(results)
        
        # Mengubah nama kolom untuk tampilan yang lebih baik
        df_recommendations.rename(columns={
            'original_location': 'Lokasi Asli Paragraf',
            'misplaced_paragraph': 'Paragraf yang Perlu Dipindah',
            'recommended_location': 'Saran Lokasi Baru'
        }, inplace=True)
        
        st.dataframe(df_recommendations, use_container_width=True)
        
        # Membuat file Excel untuk diunduh
        excel_data = create_recommendation_excel(df_recommendations)
        
        st.download_button(
            label="Unduh Laporan Rekomendasi (.xlsx)",
            data=excel_data,
            file_name=f"rekomendasi_struktur_{recommendation_file.name.split('.')[0]}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )




